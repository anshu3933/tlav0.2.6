# core/document_processing/document_loader.py
import os
from typing import List, Optional, Dict, Any, Tuple
from langchain.schema import Document
from config.logging_config import get_module_logger
from core.document_processing.document_validator import DocumentValidator

# Import specialized document loaders
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Create a logger for this module
logger = get_module_logger("document_loader")

class LoaderResult:
    """Stores the result of a document loading operation."""
    
    def __init__(self, 
                 success: bool,
                 document: Optional[Document] = None,
                 error_message: Optional[str] = None,
                 warning: Optional[str] = None):
        self.success = success
        self.document = document
        self.error_message = error_message
        self.warning = warning
    
    @property
    def has_warning(self) -> bool:
        """Check if result has a warning."""
        return self.warning is not None


class PDFLoader:
    """Handles loading and processing PDF documents with robust error handling."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load and extract text from PDF with error handling.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            A tuple of (extracted_text, error_message)
        """
        try:
            reader = PdfReader(file_path)
            
            # Check if PDF has pages
            if len(reader.pages) == 0:
                return None, "PDF file has no pages."
            
            text = ""
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {i} in {file_path}: {str(e)}")
                    # Continue with other pages
            
            # Check if we extracted any text
            if not text.strip():
                return None, "Could not extract text from PDF. The file may be scanned or contain only images."
                
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            return None, f"Error loading PDF: {str(e)}"


class DocxLoader:
    """Handles loading and processing DOCX documents with robust error handling."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load and extract text from DOCX with error handling.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            A tuple of (extracted_text, error_message)
        """
        try:
            doc = DocxDocument(file_path)
            
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(paragraphs)
            
            # Check if we extracted any text
            if not text.strip():
                return None, "Could not extract text from DOCX. The file may be empty."
                
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            return None, f"Error loading DOCX: {str(e)}"


class TextLoader:
    """Handles loading text files with robust error handling."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load text file content with error handling.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            A tuple of (file_content, error_message)
        """
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    if content.strip():
                        return content, None
                except UnicodeDecodeError:
                    continue
            
            # If we get here, none of the encodings worked
            return None, "Could not decode text file with supported encodings (utf-8, latin-1, ascii)."
                
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {str(e)}")
            return None, f"Error loading text file: {str(e)}"


class DocumentLoader:
    """Main document loading coordinator with validation and error handling."""
    
    def __init__(self):
        """Initialize with components."""
        self.validator = DocumentValidator()
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': DocxLoader(),
            '.txt': TextLoader()
        }
    
    def load_documents(self, file_paths: List[str]) -> List[LoaderResult]:
        """Load multiple documents with validation and detailed results.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of LoaderResult objects
        """
        results = []
        
        for file_path in file_paths:
            result = self.load_single_document(file_path)
            results.append(result)
                
        return results
    
    def load_single_document(self, file_path: str) -> LoaderResult:
        """Load a single document with validation and detailed result.
        
        Args:
            file_path: Path to the file
            
        Returns:
            LoaderResult object
        """
        # Validate file path
        is_valid, error_message = self.validator.validate_file_path(file_path)
        if not is_valid:
            logger.error(f"File validation failed for {file_path}: {error_message}")
            return LoaderResult(success=False, error_message=error_message)
        
        try:
            # Get file extension
            extension = os.path.splitext(file_path)[1].lower()
            
            # Get appropriate loader
            loader = self.loaders.get(extension)
            if not loader:
                error_message = f"No loader found for extension: {extension}"
                logger.error(error_message)
                return LoaderResult(success=False, error_message=error_message)
            
            # Load content
            content, load_error = loader.load(file_path)
            if load_error:
                return LoaderResult(success=False, error_message=load_error)
            
            # Validate content
            is_valid, content_error = self.validator.validate_content(content)
            if not is_valid:
                return LoaderResult(success=False, error_message=content_error)
            
            # Create document with warning if any
            warning = content_error if is_valid and content_error else None
            document = Document(
                page_content=content,
                metadata=self._create_metadata(file_path)
            )
            
            return LoaderResult(success=True, document=document, warning=warning)
            
        except Exception as e:
            logger.error(f"Unexpected error loading document {file_path}: {str(e)}", exc_info=True)
            return LoaderResult(
                success=False, 
                error_message=f"Unexpected error loading document: {str(e)}"
            )
    
    def _create_metadata(self, file_path: str) -> Dict[str, Any]:
        """Create metadata for document."""
        return {
            "source": file_path,
            "file_type": os.path.splitext(file_path)[1].lower(),
            "file_name": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "last_modified": os.path.getmtime(file_path)
        }
